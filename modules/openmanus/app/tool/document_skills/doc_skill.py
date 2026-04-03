from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .common import (
    color_to_rgb,
    ensure_image_path,
    ensure_parent_dir,
    normalize_align,
    parse_bool,
    parse_length_inch,
    parse_length_pt,
    resolve_style,
)
from .css_parser import parse_css_block_from_spec


def _to_docx_align(value: str) -> WD_ALIGN_PARAGRAPH:
    normalized = normalize_align(value, "left")
    if normalized == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if normalized == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if normalized == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _apply_run_style(run: Any, style: dict[str, Any]) -> None:
    font = run.font

    if style.get("font_name"):
        font.name = str(style["font_name"])
    if style.get("font_size") is not None:
        font.size = Pt(parse_length_pt(style.get("font_size"), 11.0))

    rgb_tuple = color_to_rgb(style.get("font_color"), (0, 0, 0))
    font.color.rgb = RGBColor(*rgb_tuple)

    font.bold = parse_bool(style.get("bold"), False)
    font.italic = parse_bool(style.get("italic"), False)
    font.underline = parse_bool(style.get("underline"), False)


def _apply_paragraph_style(paragraph: Any, style: dict[str, Any]) -> None:
    paragraph.alignment = _to_docx_align(str(style.get("align", "left")))

    fmt = paragraph.paragraph_format
    if style.get("line_height") is not None:
        fmt.line_spacing = float(style.get("line_height"))
    if style.get("space_before") is not None:
        fmt.space_before = Pt(parse_length_pt(style.get("space_before"), 0.0))
    if style.get("space_after") is not None:
        fmt.space_after = Pt(parse_length_pt(style.get("space_after"), 8.0))


def _iter_blocks(spec: dict[str, Any]) -> list[dict[str, Any]]:
    sections = spec.get("sections")
    if isinstance(sections, list):
        return [section for section in sections if isinstance(section, dict)]

    blocks = spec.get("blocks")
    if isinstance(blocks, list):
        return [block for block in blocks if isinstance(block, dict)]

    content = spec.get("content")
    if isinstance(content, list):
        return [{"type": "paragraph", "text": str(item)} for item in content]

    return [
        {
            "type": "heading",
            "text": str(spec.get("title", "Document")),
            "level": 1,
            "style_ref": "title",
        },
        {
            "type": "paragraph",
            "text": str(spec.get("content", "")),
            "style_ref": "body",
        },
    ]


def _add_text_paragraph(document: Any, text: str, style: dict[str, Any], list_level: int | None = None) -> None:
    paragraph = document.add_paragraph()
    if list_level is not None and list_level >= 0:
        paragraph.style = "List Bullet"

    run = paragraph.add_run(text)
    _apply_run_style(run, style)
    _apply_paragraph_style(paragraph, style)


def create_docx_from_spec(spec: dict[str, Any], output_path: str) -> str:
    document = Document()

    default_style = spec.get("default_style") if isinstance(spec.get("default_style"), dict) else {}
    css_styles = parse_css_block_from_spec(spec)

    page = spec.get("page") if isinstance(spec.get("page"), dict) else {}
    if page:
        section = document.sections[0]
        section.left_margin = Inches(parse_length_inch(page.get("margin_left"), 0.8))
        section.right_margin = Inches(parse_length_inch(page.get("margin_right"), 0.8))
        section.top_margin = Inches(parse_length_inch(page.get("margin_top"), 0.8))
        section.bottom_margin = Inches(parse_length_inch(page.get("margin_bottom"), 0.8))

    for block in _iter_blocks(spec):
        block_type = str(block.get("type", "paragraph")).lower()

        style = resolve_style(
            default_style=default_style,
            css_styles=css_styles,
            style_ref=block.get("style_ref"),
            inline_style=block.get("style") if isinstance(block.get("style"), dict) else None,
        )

        if block_type in {"heading", "title"}:
            level = int(block.get("level", 1) or 1)
            level = min(max(level, 1), 9)
            paragraph = document.add_heading(level=level)
            run = paragraph.add_run(str(block.get("text", "")))
            _apply_run_style(run, style)
            _apply_paragraph_style(paragraph, style)
            continue

        if block_type in {"paragraph", "text"}:
            _add_text_paragraph(document, str(block.get("text", "")), style)
            continue

        if block_type in {"list", "bullet-list", "bullet"}:
            items = block.get("items") if isinstance(block.get("items"), list) else []
            for item in items:
                _add_text_paragraph(document, str(item), style, list_level=0)
            continue

        if block_type in {"image", "picture"} and block.get("path"):
            image_path = ensure_image_path(str(block["path"]))
            width = block.get("width")
            document.add_picture(
                str(image_path),
                width=Inches(parse_length_inch(width, 4.0)) if width is not None else None,
            )
            continue

        if block_type == "table":
            rows = block.get("rows") if isinstance(block.get("rows"), list) else []
            if not rows:
                continue

            normalized_rows = []
            for row in rows:
                if isinstance(row, list):
                    normalized_rows.append([str(cell) for cell in row])

            if not normalized_rows:
                continue

            max_cols = max(len(row) for row in normalized_rows)
            table = document.add_table(rows=len(normalized_rows), cols=max_cols)
            style_name = block.get("table_style") or block.get("style_name")
            if isinstance(style_name, str) and style_name.strip():
                table.style = style_name

            for row_index, row_values in enumerate(normalized_rows):
                for col_index, cell_value in enumerate(row_values):
                    paragraph = table.cell(row_index, col_index).paragraphs[0]
                    run = paragraph.add_run(cell_value)
                    _apply_run_style(run, style)
                    _apply_paragraph_style(paragraph, style)

    destination = ensure_parent_dir(output_path)
    document.save(str(destination))
    return str(destination)
